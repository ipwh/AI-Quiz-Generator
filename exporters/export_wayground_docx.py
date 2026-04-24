# exporters/export_wayground_docx.py

from io import BytesIO
from docx import Document

_CORRECT_MAP = {"1": "A", "2": "B", "3": "C", "4": "D"}

def _to_letter(correct_val) -> str:
    """將 correct 欄位（1~4、"1"~"4"、list）轉為 A~D。"""
    if isinstance(correct_val, list):
        correct_val = correct_val[0] if correct_val else "1"
    val = str(correct_val).strip().split(".")[0]  # 處理 "1.0"
    return _CORRECT_MAP.get(val, val)


def export_wayground_docx(df, subject: str) -> bytes:
    """
    匯出 Wayground / 教學用 DOCX。
    正確答案顯示為 A~D，而非 1~4。
    """
    doc = Document()
    doc.add_heading(f"{subject} 題目練習", level=1)

    for i, r in df.iterrows():
        doc.add_paragraph(f"{i + 1}. {r['question']}")

        doc.add_paragraph(f"A. {r['option_1']}")
        doc.add_paragraph(f"B. {r['option_2']}")
        doc.add_paragraph(f"C. {r['option_3']}")
        doc.add_paragraph(f"D. {r['option_4']}")

        answer_letter = _to_letter(r["correct"])
        doc.add_paragraph(f"✅ 正確答案：{answer_letter}")

        if r.get("explanation"):
            doc.add_paragraph(f"解說：{r['explanation']}")

        doc.add_paragraph("")

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
